"""Render a Markdown subset to a .docx byte stream (python-docx, lazy import).

Supports the elements an agent produces for a synthesis document: ATX headings
(``#``..``######``), paragraphs, bullet and numbered lists, GitHub pipe tables,
and inline ``**bold**`` / ``*italic*``. Anything fancier degrades to plain text.

CPU bound and synchronous; run it off the event loop (``asyncio.to_thread``).
"""

from __future__ import annotations

import io
import re

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_ITALIC = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
_BULLET = re.compile(r"^\s*[-*+]\s+(.*)$")
_NUMBERED = re.compile(r"^\s*\d+[.)]\s+(.*)$")
_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_TABLE_ROW = re.compile(r"^\s*\|(.+)\|\s*$")
_TABLE_SEP = re.compile(r"^\s*\|?[\s:|-]+\|?\s*$")


def _add_inline(paragraph: object, text: str) -> None:
    """Append ``text`` to a paragraph, honouring **bold** and *italic* runs."""
    pos = 0
    # Bold first, then italic within the non-bold gaps, kept simple and robust.
    for match in _BOLD.finditer(text):
        if match.start() > pos:
            _add_italic(paragraph, text[pos : match.start()])
        run = paragraph.add_run(match.group(1))  # type: ignore[attr-defined]
        run.bold = True
        pos = match.end()
    if pos < len(text):
        _add_italic(paragraph, text[pos:])


def _add_italic(paragraph: object, text: str) -> None:
    pos = 0
    for match in _ITALIC.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])  # type: ignore[attr-defined]
        run = paragraph.add_run(match.group(1))  # type: ignore[attr-defined]
        run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])  # type: ignore[attr-defined]


def _split_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def markdown_to_docx(markdown: str, *, title: str | None = None) -> bytes:
    """Convert a Markdown subset to .docx bytes."""
    import docx  # noqa: PLC0415

    document = docx.Document()
    if title:
        document.add_heading(title, level=0)

    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue

        heading = _HEADING.match(line)
        if heading:
            document.add_heading(heading.group(2).strip(), level=min(len(heading.group(1)), 6))
            index += 1
            continue

        # A pipe table: a header row, a separator row, then body rows.
        if _TABLE_ROW.match(line) and index + 1 < len(lines) and _TABLE_SEP.match(lines[index + 1]):
            header = _split_row(line)
            body: list[list[str]] = []
            index += 2
            while index < len(lines) and _TABLE_ROW.match(lines[index]):
                body.append(_split_row(lines[index]))
                index += 1
            table = document.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for col, cell_text in enumerate(header):
                _add_inline(table.rows[0].cells[col].paragraphs[0], cell_text)
            for row_cells in body:
                cells = table.add_row().cells
                for col in range(len(header)):
                    value = row_cells[col] if col < len(row_cells) else ""
                    _add_inline(cells[col].paragraphs[0], value)
            continue

        bullet = _BULLET.match(line)
        if bullet:
            _add_inline(document.add_paragraph(style="List Bullet"), bullet.group(1).strip())
            index += 1
            continue

        numbered = _NUMBERED.match(line)
        if numbered:
            _add_inline(document.add_paragraph(style="List Number"), numbered.group(1).strip())
            index += 1
            continue

        _add_inline(document.add_paragraph(), stripped)
        index += 1

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
